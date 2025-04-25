from docx import Document
from datetime import datetime #datetime é uma biblioteca usada para puxar a data atual
import pandas as pd

tabela_excel = pd.read_excel("Planilha.xlsx").fillna("") #.fillna("") é usado para preencher células vazias com string vazia

for linha in tabela_excel.index:
    documento = Document("Termo.docx")

    nome = tabela_excel.loc[linha, "Nome"] #Pega os dados de cada coluna por linha
    cpf = tabela_excel.loc[linha, "CPF"]
    modelo = tabela_excel.loc[linha, "Modelo"]
    st = tabela_excel.loc[linha, "Service Tag"]
    cidade = tabela_excel.loc[linha, "Cidade"]
    empresa = tabela_excel.loc[linha, "Empresa"]
    marca = tabela_excel.loc[linha, "Marca"]


    referencias = { #Cria um dicionário com os códigos do documento Word e os valores a serem substituídos
        "XXXX": str(nome),
        "ZZZZ": str(cpf),  #Usando str() para garantir que venha como um texto
        "WWWW": str(modelo),
        "AAAA": str(st),
        "BBBB": str(marca),
        "CCCC": str(cidade),
        "EEEE": str(empresa),
        "DD": datetime.now().strftime("%d"), #.strftime("") é usado para formatar a data corretamente
        "MM": datetime.now().strftime("%m"), #Se formatar como string normalmente, o dia/mês pode vir como "4" ao invés de 04
        "YYYY": datetime.now().strftime("%Y"),
    }

    def substituir_em_runs(runs): #Função usada para manter a formatação do texto, negrito, itálico... basicamente substituir os códigos dentro dos "runs" (trechos de texto) do Word
        for run in runs:
            for codigo, valor in referencias.items():
                if codigo in run.text:
                    run.text = run.text.replace(codigo, valor)


    for paragrafo in documento.paragraphs:  #Substituir nos parágrafos
        substituir_em_runs(paragrafo.runs)

    
    for tabela_word in documento.tables:  #Faz as substituições também nas tabelas do documento word
        for linha_tabela in tabela_word.rows:
            for celula in linha_tabela.cells:
                for paragrafo in celula.paragraphs:
                    substituir_em_runs(paragrafo.runs)

    documento.save(f"Termo - {nome}.docx") #Salvando o documento para cada pessoa da planilha e com o seu nome
